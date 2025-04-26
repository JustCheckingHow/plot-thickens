from docx import Document
import re
from typing import List, Tuple
import os
import fitz  # PyMuPDF


def docx_to_markdown_chapters(file_path: str) -> tuple[List[str], List[str]]:
    """
    Converts a Word document to a list of markdown chapters based on the highest level heading.

    Args:
        file_path: Path to the Word document file

    Returns:
        Tuple containing:
        - List of markdown strings, each representing a chapter
        - List of chapter names extracted from highest level headings
    """
    # Load the document
    doc = Document(file_path)

    # Extract headings and their levels from the document
    headings = []
    for paragraph in doc.paragraphs:
        if paragraph.style.name.startswith("Heading"):
            level = int(paragraph.style.name.split(" ")[1])
            headings.append((level, paragraph.text))

    # If no headings found, return the whole document as a single chapter
    if not headings:
        return [_convert_document_to_markdown(doc)], ["Untitled"]

    # Find the highest level heading (lowest number)
    min_level = min([level for level, _ in headings]) if headings else 0

    # Split document into chapters based on highest level headings
    chapter_indices = []
    for i, para in enumerate(doc.paragraphs):
        if para.style.name.startswith("Heading"):
            level = int(para.style.name.split(" ")[1])
            if level == min_level:
                chapter_indices.append(i)

    # Add a final index for the end of document
    chapter_indices.append(len(doc.paragraphs))

    # Extract chapters and their names
    chapters = []
    chapter_names = []
    for i in range(len(chapter_indices) - 1):
        start_idx = chapter_indices[i]
        end_idx = chapter_indices[i + 1]
        chapter_paras = doc.paragraphs[start_idx:end_idx]
        chapters.append(_convert_paragraphs_to_markdown(chapter_paras))

        # Extract chapter name from the heading
        if start_idx < len(doc.paragraphs):
            chapter_names.append(doc.paragraphs[start_idx].text)

    return chapters, chapter_names


def _convert_paragraphs_to_markdown(paragraphs) -> str:
    """
    Converts a list of paragraphs to markdown format, preserving heading levels.

    Args:
        paragraphs: List of docx Paragraph objects

    Returns:
        Markdown formatted string
    """
    markdown = []

    for para in paragraphs:
        if para.style.name.startswith("Heading"):
            level = int(para.style.name.split(" ")[1])
            markdown.append(f"{'#' * level} {para.text}")
        elif para.style.name == "Title":
            markdown.append(f"# {para.text}")
        else:
            # Add paragraph text with proper new lines
            if para.text.strip():
                markdown.append(para.text)

    return "\n\n".join(markdown)


def _convert_document_to_markdown(doc) -> str:
    """
    Converts an entire document to markdown format.

    Args:
        doc: Document object

    Returns:
        Markdown formatted string
    """
    return _convert_paragraphs_to_markdown(doc.paragraphs)


def save_chapters_to_files(
    chapters: List[str], output_dir: str, base_filename: str = "chapter"
) -> List[str]:
    """
    Saves chapters to individual markdown files.

    Args:
        chapters: List of markdown formatted chapter strings
        output_dir: Directory to save the files
        base_filename: Base name for the files

    Returns:
        List of saved file paths
    """
    os.makedirs(output_dir, exist_ok=True)

    file_paths = []
    for i, chapter in enumerate(chapters):
        file_path = os.path.join(output_dir, f"{base_filename}_{i+1}.md")
        with open(file_path, "w", encoding="utf-8") as f:
            f.write(chapter)
        file_paths.append(file_path)

    return file_paths


def pdf_to_markdown_chapters(file_path: str) -> Tuple[List[str], List[str]]:
    """
    Converts a PDF document to a list of markdown chapters based on the table of contents.

    Args:
        file_path: Path to the PDF document file

    Returns:
        Tuple containing:
        - List of markdown strings, each representing a chapter
        - List of chapter names extracted from the table of contents
    """
    # Load the document
    doc = fitz.open(file_path)

    # Extract table of contents
    toc = doc.get_toc()

    # If no TOC found, return the whole document as a single chapter
    if not toc:
        return [_convert_pdf_to_markdown(doc)], ["Untitled"]

    # Filter for top-level entries (usually chapters)
    # TOC entries are [level, title, page, ...]
    top_level = [entry for entry in toc if entry[0] == 1]

    # If no top-level entries, try using all entries
    if not top_level:
        top_level = toc

    # Extract chapter page ranges
    chapter_ranges = []
    for i in range(len(top_level)):
        start_page = top_level[i][2] - 1  # Convert from 1-indexed to 0-indexed
        end_page = top_level[i + 1][2] - 2 if i < len(top_level) - 1 else len(doc) - 1
        chapter_title = top_level[i][1]
        chapter_ranges.append((start_page, end_page, chapter_title))

    # Convert chapters to markdown
    chapters = []
    chapter_names = []

    for start_page, end_page, title in chapter_ranges:
        # Ensure start_page is not negative
        start_page = max(0, start_page)
        # Ensure end_page is not beyond document length
        end_page = min(end_page, len(doc) - 1)
        # Ensure start_page doesn't exceed end_page
        if start_page <= end_page:
            chapter_text = _extract_text_from_page_range(doc, start_page, end_page)
            chapters.append(chapter_text)
            chapter_names.append(title)

    return chapters, chapter_names


def _extract_text_from_page_range(doc, start_page: int, end_page: int) -> str:
    """
    Extracts text from a range of pages and converts to markdown.

    Args:
        doc: PDF document
        start_page: Starting page number
        end_page: Ending page number

    Returns:
        Markdown formatted string
    """
    text_blocks = []

    for page_num in range(start_page, end_page + 1):
        page = doc[page_num]
        text = page.get_text()
        if text.strip():
            text_blocks.append(text)

    return "\n\n".join(text_blocks)


def _convert_pdf_to_markdown(doc) -> str:
    """
    Converts an entire PDF document to markdown format.

    Args:
        doc: PDF document

    Returns:
        Markdown formatted string
    """
    return _extract_text_from_page_range(doc, 0, len(doc) - 1)


if __name__ == "__main__":
    # chapters, chapter_names = docx_to_markdown_chapters(
    #     "/Users/jm/repos/plot-thickens/data/Book with inconsistency.docx"
    # )
    # save_chapters_to_files(chapters, "output")
    chapters, chapter_names = pdf_to_markdown_chapters(
        "/Users/jm/repos/plot-thickens/data/Fifty-Shades-of-Grey-PDF.pdf"
    )
    # print(chapters)
    print(chapter_names)
