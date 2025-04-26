from pydantic import BaseModel
import re
import io
from typing import Optional
from docx import Document


class MarkdownToWordRequest(BaseModel):
    markdown_content: str
    filename: Optional[str] = "document"
    author_name: Optional[str] = "PlotDickens"


async def markdown_to_docx_with_comments(request: MarkdownToWordRequest) -> bytes:
    """Converts markdown content to a Microsoft Word document with proper comments"""

    # Regular expression to find comments in the text <comment id=...></comment>
    comment_pattern = re.compile(r"(.*?)<comment id=(.*?)>(.*?)</comment>(.*?)(?=<comment|\Z)", re.DOTALL)
    # Regular expression to identify markdown headings
    heading_pattern = re.compile(r"^(#{1,6})\s+(.+?)$", re.MULTILINE)

    # Create a new Word document
    doc = Document()

    # Process the content and extract comments
    remaining_text = request.markdown_content
    
    # Find all comments in the text
    matches = list(comment_pattern.finditer(remaining_text))
    
    if not matches:
        # No comments in text, just process the markdown normally
        lines = remaining_text.split("\n")
        for line in lines:
            heading_match = heading_pattern.match(line.strip())
            if heading_match:
                level = len(heading_match.group(1))
                heading_text = heading_match.group(2).strip()
                paragraph = doc.add_paragraph(heading_text)
                paragraph.style = f"Heading {level}"
            else:
                line = line.strip()
                if line:
                    doc.add_paragraph(line)
    else:
        # Process each comment section
        for match in matches:
            text_before_comment = match.group(1)
            comment_id = match.group(2)
            comment_text = match.group(3)
            text_after_comment = match.group(4)
            
            # Process text before comment
            if text_before_comment:
                lines = text_before_comment.split("\n")
                for i, line in enumerate(lines):
                    heading_match = heading_pattern.match(line.strip())
                    if heading_match:
                        level = len(heading_match.group(1))
                        heading_text = heading_match.group(2).strip()
                        paragraph = doc.add_paragraph(heading_text)
                        paragraph.style = f"Heading {level}"
                    else:
                        line = line.strip()
                        if line:
                            # For the last line, we'll add the comment inline
                            if i == len(lines) - 1:
                                paragraph = doc.add_paragraph(line)
                                # Add the comment directly
                                run = paragraph.add_run(" ")  # Space after text
                                run.add_comment(comment_text, author=request.author_name, initials="A")
                                
                                # If there's text after the comment on the same line, add it to the same paragraph
                                if text_after_comment and text_after_comment.strip() and not text_after_comment.startswith("\n"):
                                    first_line_after = text_after_comment.split("\n")[0].strip()
                                    if first_line_after:
                                        paragraph.add_run(first_line_after)
                                        
                                    # Process the rest of the text after the comment if it continues to new lines
                                    if "\n" in text_after_comment:
                                        rest_lines = text_after_comment.split("\n")[1:]
                                        for rest_line in rest_lines:
                                            rest_line = rest_line.strip()
                                            if rest_line:
                                                doc.add_paragraph(rest_line)
                            else:
                                doc.add_paragraph(line)
            else:
                # Comment is at the beginning of a line, add it to a new paragraph
                paragraph = doc.add_paragraph()
                run = paragraph.add_run("")  # Empty run
                run.add_comment(comment_text, author=request.author_name, initials="A")
                
                # If there's text after the comment, add it to the same paragraph
                if text_after_comment and text_after_comment.strip() and not text_after_comment.startswith("\n"):
                    first_line_after = text_after_comment.split("\n")[0].strip()
                    if first_line_after:
                        paragraph.add_run(first_line_after)
                    
                    # Process the rest of the text after the comment if it continues to new lines
                    if "\n" in text_after_comment:
                        rest_lines = text_after_comment.split("\n")[1:]
                        for rest_line in rest_lines:
                            rest_line = rest_line.strip()
                            if rest_line:
                                doc.add_paragraph(rest_line)
        
        # Process any remaining text after the last comment
        if matches[-1].end() < len(remaining_text):
            remaining = remaining_text[matches[-1].end():]
            lines = remaining.split("\n")
            for line in lines:
                heading_match = heading_pattern.match(line.strip())
                if heading_match:
                    level = len(heading_match.group(1))
                    heading_text = heading_match.group(2).strip()
                    paragraph = doc.add_paragraph(heading_text)
                    paragraph.style = f"Heading {level}"
                else:
                    line = line.strip()
                    if line:
                        doc.add_paragraph(line)
                        
    # Save the document to a bytes buffer
    docx_buffer = io.BytesIO()
    doc.save(docx_buffer)
    docx_buffer.seek(0)

    # Set up the response with appropriate headers
    filename = f"{request.filename}.docx"
    headers = {
        "Content-Disposition": f'attachment; filename="{filename}"',
        "Content-Type": "application/vnd.openxmlformats-officedocument.wordprocessingml.document",
    }

    return docx_buffer.getvalue()


if __name__ == "__main__":
    import asyncio

    # Example markdown content with comments
    markdown_content = """# Sample Document
    
    This is a sample document with some text.
    
    Here is a paragraph with a comment <comment id="123">This is a comment about something important</comment>.
    
    And another paragraph with no comments.
    
    <comment id="456">Another comment</comment> at the beginning of a line.
    """

    # Create the request object
    request = MarkdownToWordRequest(
        markdown_content=markdown_content, filename="example_document"
    )

    # Run the async function and save the result
    async def main():
        docx_bytes = await markdown_to_docx_with_comments(request)

        # Save to a file
        with open("example_document.docx", "wb") as f:
            f.write(docx_bytes)

        print("Document saved as 'example_document.docx'")

    asyncio.run(main())
